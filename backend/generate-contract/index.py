import json
import os
import io
import base64
from typing import Dict, Any
from docx import Document
from docx.shared import Inches
import psycopg2
import requests
from PIL import Image

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    '''
    Business: Generate license agreement DOCX from template and send to Telegram
    Args: event - dict with httpMethod, body containing form data
          context - object with request_id attribute
    Returns: HTTP response dict with success/error status
    '''
    method: str = event.get('httpMethod', 'GET')
    
    if method == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Max-Age': '86400'
            },
            'body': ''
        }
    
    if method != 'POST':
        return {
            'statusCode': 405,
            'headers': {'Access-Control-Allow-Origin': '*', 'Content-Type': 'application/json'},
            'body': json.dumps({'error': 'Method not allowed'})
        }
    
    try:
        body_data = json.loads(event.get('body', '{}'))
        print(f'Received payload: {json.dumps(body_data, ensure_ascii=False)}')
        
        db_url = os.environ.get('DATABASE_URL')
        if not db_url:
            raise Exception('DATABASE_URL not found')
        
        conn = psycopg2.connect(db_url)
        cur = conn.cursor()
        
        cur.execute('UPDATE contract_counter SET current_number = current_number + 1, updated_at = CURRENT_TIMESTAMP WHERE id = 1 RETURNING current_number')
        result = cur.fetchone()
        contract_number = result[0] if result else 1
        
        conn.commit()
        
        cur.execute('SELECT template_data FROM template_storage WHERE id = 1')
        template_result = cur.fetchone()
        
        cur.close()
        conn.close()
        
        if not template_result or not template_result[0]:
            raise Exception('Template not found. Please upload template.docx first via /434 page')
        
        template_data = template_result[0]
        if isinstance(template_data, memoryview):
            template_bytes = template_data.tobytes()
        else:
            template_bytes = bytes(template_data)
        
        print(f'Template size: {len(template_bytes)} bytes')
        print(f'First 20 bytes (hex): {template_bytes[:20].hex()}')
        print(f'First 4 bytes should be 504B0304 (PK..): {template_bytes[:4].hex()}')
        
        if not template_bytes.startswith(b'PK'):
            raise Exception(f'Template is corrupted. First bytes: {template_bytes[:10].hex()}. Please re-upload template.docx via /434')
        
        template_io = io.BytesIO(template_bytes)
        
        replacements = {
            '{{номер_договора}}': str(contract_number),
            '{{дата_заключения_договора}}': body_data.get('дата_заключения_договора', ''),
            '{{graj}}': body_data.get('graj', ''),
            '{{ФИО_ИП_полностью_кого}}': body_data.get('ФИО_ИП_полностью_кого', ''),
            '{{ФИО_ИП_кратко}}': body_data.get('ФИО_ИП_кратко', ''),
            '{{NIK}}': body_data.get('NIK', ''),
            '{{PAS}}': body_data.get('PAS', ''),
            '{{mail}}': body_data.get('mail', ''),
            '{{ИНН_SWIFT}}': body_data.get('ИНН_SWIFT', ''),
            '{{РЕКВИЗИТЫ_БАНК}}': body_data.get('РЕКВИЗИТЫ_БАНК', ''),
            '{{naz}}': body_data.get('naz', ''),
            '{{isp}}': body_data.get('isp', ''),
            '{{avt}}': body_data.get('avt', ''),
            '{{avttext}}': body_data.get('avttext', ''),
            '{{fongr}}': body_data.get('fongr', ''),
            '{{procc}}': body_data.get('procc', '')
        }
        
        print(f'Replacements map:')
        for key, value in replacements.items():
            print(f'  {key} => {value}')
        
        doc = Document(template_io)
        
        cover_image_bytes = None
        if body_data.get('cover_image'):
            try:
                import base64 as b64
                img_data = b64.b64decode(body_data['cover_image'])
                img = Image.open(io.BytesIO(img_data))
                img = img.convert('RGB')
                img_resized = img.resize((150, 150), Image.LANCZOS)
                img_io = io.BytesIO()
                img_resized.save(img_io, format='PNG')
                cover_image_bytes = img_io.getvalue()
                print('Cover image resized to 150x150')
            except Exception as e:
                print(f'Failed to process cover image: {e}')
        
        def replace_in_paragraph(paragraph):
            """Replace placeholders in paragraph (handles runs splitting)"""
            full_text = paragraph.text
            
            if '{{img}}' in full_text and cover_image_bytes:
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    img_stream = io.BytesIO(cover_image_bytes)
                    paragraph.runs[0].add_picture(img_stream, width=Inches(1.57))
                return
            
            has_replacements = False
            for key, value in replacements.items():
                if key in full_text:
                    full_text = full_text.replace(key, value)
                    has_replacements = True
            
            if has_replacements:
                first_run_font = paragraph.runs[0].font if paragraph.runs else None
                
                for run in paragraph.runs:
                    run.text = ''
                
                if paragraph.runs and first_run_font:
                    paragraph.runs[0].text = full_text
                    paragraph.runs[0].font.bold = first_run_font.bold
                    paragraph.runs[0].font.italic = first_run_font.italic
                    paragraph.runs[0].font.underline = first_run_font.underline
                    paragraph.runs[0].font.size = first_run_font.size
                    paragraph.runs[0].font.name = first_run_font.name
                elif paragraph.runs:
                    paragraph.runs[0].text = full_text
        
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph)
        
        for section in doc.sections:
            for paragraph in section.header.paragraphs:
                replace_in_paragraph(paragraph)
            for paragraph in section.footer.paragraphs:
                replace_in_paragraph(paragraph)
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        telegram_token = os.environ.get('TELEGRAM_BOT_TOKEN', '')
        chat_id = os.environ.get('TELEGRAM_CHAT_ID', '')
        
        nickname = body_data.get('NIK', 'Unknown')
        
        cover_image_b64 = body_data.get('cover_image', '')
        if cover_image_b64:
            import base64
            cover_image_bytes = base64.b64decode(cover_image_b64)
            cover_image_name = body_data.get('cover_image_name', 'cover.jpg')
            
            photo_files = {'photo': (cover_image_name, io.BytesIO(cover_image_bytes))}
            photo_data = {
                'chat_id': chat_id,
                'caption': f'🎨 Обложка для {nickname}'
            }
            photo_url = f'https://api.telegram.org/bot{telegram_token}/sendPhoto'
            photo_response = requests.post(photo_url, files=photo_files, data=photo_data)
            
            if photo_response.status_code != 200:
                raise Exception(f'Telegram photo error: {photo_response.text}')
        
        files = {
            'document': (f'{nickname}_Договор_{contract_number}.docx', output, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        }
        data = {
            'chat_id': chat_id,
            'caption': f'🎵 {nickname} - Лицензионный договор №{contract_number}'
        }
        
        telegram_url = f'https://api.telegram.org/bot{telegram_token}/sendDocument'
        response = requests.post(telegram_url, files=files, data=data)
        
        if response.status_code != 200:
            raise Exception(f'Telegram API error: {response.text}')
        
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Content-Type': 'application/json'
            },
            'isBase64Encoded': False,
            'body': json.dumps({
                'success': True,
                'contract_number': contract_number,
                'message': 'Договор успешно сгенерирован и отправлен в Telegram'
            })
        }
        
    except Exception as e:
        error_msg = str(e)
        print(f'ERROR in generate-contract: {error_msg}')
        import traceback
        traceback.print_exc()
        
        return {
            'statusCode': 500,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Content-Type': 'application/json'
            },
            'isBase64Encoded': False,
            'body': json.dumps({
                'success': False,
                'error': error_msg
            })
        }